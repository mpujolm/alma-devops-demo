import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, '2E4057')
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = 'F0F4F8' if ri % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            # Monospace for API names (first col usually)
            if i == 0 and val.endswith('__c') or (i == 0 and val.startswith('`')):
                run.font.name = 'Courier New'
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# --- Title ---
title = doc.add_heading('Data Dictionary — Alma Natural Cosmetics', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph('This document describes the custom data model of the Alma Natural Cosmetics Salesforce org.')
doc.add_paragraph()

# --- Relationship diagram (text) ---
doc.add_heading('Object Relationships', 1)
p = doc.add_paragraph()
p.style = 'No Spacing'
run = p.add_run(
    'Contact  ──────────────────────────  (Lookup)\n'
    '  └─►  OrderRequest__c  ────────────  (Master)\n'
    '         └─►  OrderRequestLineItem__c  (Detail)\n'
    '                └─►  Product2  ───────  (Lookup)\n'
    '                       └─►  Recipe__c  ─  (Lookup)\n'
    '                              └─►  Recipe_Line_Item__c  (Lookup)\n'
    '                                     └─►  Raw_Material__c  (Lookup)'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)
doc.add_paragraph()

# ============================================================
# 1. CONTACT
# ============================================================
doc.add_heading('Contact  (standard object — custom fields)', 1)
doc.add_paragraph('Standard Salesforce object representing a customer.')
add_table(doc,
    ['API Name', 'Label', 'Type', 'Description'],
    [
        ['Customer_Type__c',           'Customer Type',           'Picklist',            'Segmentation: Family, Friend, School, Known, New, Other'],
        ['Communication_Preferences__c','Communication Preferences','Multiselect Picklist','Preferred channels: Email, Phone, WhatsApp, In Person'],
        ['Opt_In__c',                  'Opt In',                  'Checkbox',            'Marketing communications consent'],
        ['Opt_In_Date__c',             'Opt In Date',             'Date',                'Date consent was given'],
        ['Product_Categories__c',      'Product Categories',      'Multiselect Picklist','Categories of interest: Cleaning, Cosmetics, Baby Cosmetics, Gels, Shampoo'],
        ['Preference_Notes__c',        'Preference Notes',        'Long Text Area',      'Preferences, special requests, personal notes'],
        ['First_Order_Request__c',     'First Order Request',     'Date',                "Date of the customer's first order"],
        ['Last_Order_Request__c',      'Last Order Request',      'Date',                "Date of the customer's most recent order"],
    ],
    col_widths=[4.5, 4, 3.5, 6.5]
)
doc.add_paragraph()

# ============================================================
# 2. ORDERREQUEST__C
# ============================================================
doc.add_heading('OrderRequest__c — Order', 1)
doc.add_paragraph('Stores customer order headers. Auto-number: ORD-{0000}. Sharing model: ReadWrite.')
add_table(doc,
    ['API Name', 'Label', 'Type', 'Req', 'Description'],
    [
        ['Name',               'Order Number',  'Auto-number',                       '—', 'Auto-generated identifier, e.g. ORD-0001'],
        ['Contact__c',         'Contact',       'Lookup → Contact',                  '—', 'Customer who placed the order'],
        ['Order_Date__c',      'Order Date',    'Date',                              '✓', 'Date the order was placed'],
        ['Status__c',          'Status',        'Picklist',                          '✓', 'Requested · In Progress · Ready · Delivered'],
        ['Payment_Method__c',  'Payment Method','Picklist',                          '—', 'Bizum · Cash · Bank Transfer'],
        ['Total_Amount__c',    'Total Amount',  'Currency',                          '—', 'Calculated by Flow — sum of all line items including deposits'],
        ['Notes__c',           'Notes',         'Long Text Area',                    '—', 'Internal notes, delivery instructions, special requests'],
    ],
    col_widths=[4.5, 3.5, 4, 1, 5.5]
)
doc.add_paragraph()

# ============================================================
# 3. ORDERREQUESTLINEITEM__C
# ============================================================
doc.add_heading('OrderRequestLineItem__c — Order Line Item', 1)
doc.add_paragraph('Individual product lines within an order. Auto-number: LI-{0000}. Sharing model: Controlled by Parent.')
add_table(doc,
    ['API Name', 'Label', 'Type', 'Req', 'Description'],
    [
        ['Name',                 'Line Item Name',  'Auto-number',                            '—', 'Auto-generated identifier, e.g. LI-0001'],
        ['OrderRequest__c',      'Order Request',   'Master-Detail → OrderRequest__c',        '✓', 'Parent order — deleting the order deletes all its line items'],
        ['Product__c',           'Product',         'Lookup → Product2',                      '—', 'Product from the Alma catalogue'],
        ['Quantity__c',          'Quantity',        'Number',                                 '✓', 'Units ordered'],
        ['Unit_Price__c',        'Unit Price',      'Currency',                               '✓', 'Price per unit at the time of the order'],
        ['Deposit_Container__c', 'Deposit Container','Checkbox',                              '—', 'Whether a container deposit is included (return programme)'],
        ['Deposit_Price__c',     'Deposit Price',   'Currency',                               '—', 'Deposit amount — discounted on next order when container is returned'],
        ['Photo__c',             'Photo',           'URL',                                    '—', 'Link to product photo (Google Drive or external storage)'],
    ],
    col_widths=[4.5, 3.5, 4, 1, 5.5]
)
doc.add_paragraph()

# ============================================================
# 4. PRODUCT2
# ============================================================
doc.add_heading('Product2 — Product  (standard object — custom fields)', 1)
doc.add_paragraph('Standard Salesforce product catalogue, extended with Alma-specific fields.')
add_table(doc,
    ['API Name', 'Label', 'Type', 'Description'],
    [
        ['Price__c',          'Price',           'Currency', 'Selling price'],
        ['Production_Cost__c','Production Cost', 'Currency', 'Production cost, excluding container deposit'],
        ['Deposit_Price__c',  'Deposit Price',   'Currency', 'Container deposit amount for this product'],
        ['Stock__c',          'Stock',           'Number',   'Units currently available'],
        ['Batch_Production__c','Batch Production','Number',  'Units produced per batch'],
        ['Quantity_ml__c',    'Quantity (ml)',   'Number',   'Product volume in ml'],
        ['Euros_Margin__c',   'Euros Margin',    'Currency', 'Margin in euros (Price − Production Cost)'],
        ['Margin_Percent__c', 'Margin %',        'Percent',  '(Margin / Cost) × 100'],
        ['Recipe__c',         'Recipe',          'Lookup → Recipe__c', 'Recipe linked to this product'],
        ['Product_Image__c',  'Product Image',   'HTML',     'Image of the finished product'],
    ],
    col_widths=[4.5, 4, 3.5, 6.5]
)
doc.add_paragraph()

# ============================================================
# 5. RECIPE__C
# ============================================================
doc.add_heading('Recipe__c — Cosmetic Recipe', 1)
doc.add_paragraph('Stores the formulation (recipe) for each cosmetic product.')
add_table(doc,
    ['API Name', 'Label', 'Type', 'Description'],
    [
        ['Name',                'Recipe Name',       'Text',    'Name of the recipe'],
        ['Description__c',      'Description',       'Text',    'Recipe description'],
        ['Product_Category__c', 'Product Category',  'Picklist','Cleaning product · Cosmetic · Kids cosmetic · Shower gel · Shampoo'],
        ['Skin_Type__c',        'Skin Type',         'Picklist','Normal · Dry Skin · Oily Skin · Family Skin Care'],
        ['PAO__c',              'PAO',               'Number',  'Period After Opening — months of use once opened'],
        ['Expiration_Date__c',  'Expiration Date',   'Number',  'Shelf life in months (unopened)'],
        ['Total_Quantity_ml__c','Total Quantity (ml)','Number', 'Total ml produced by the recipe'],
    ],
    col_widths=[4.5, 4, 3.5, 6.5]
)
doc.add_paragraph()

# ============================================================
# 6. RECIPE_LINE_ITEM__C
# ============================================================
doc.add_heading('Recipe_Line_Item__c — Recipe Ingredient Line', 1)
doc.add_paragraph('Each row represents one ingredient (raw material) and its quantity within a recipe.')
add_table(doc,
    ['API Name', 'Label', 'Type', 'Description'],
    [
        ['Recipe__c',      'Recipe',       'Lookup → Recipe__c',       'Parent recipe'],
        ['Raw_Material__c','Raw Material', 'Lookup → Raw_Material__c', 'Ingredient used'],
        ['Quantity__c',    'Quantity',     'Number',                   'Amount in grams'],
        ['Price__c',       'Price',        'Currency',                 'Cost of this ingredient in the recipe'],
    ],
    col_widths=[4.5, 4, 3.5, 6.5]
)
doc.add_paragraph()

# ============================================================
# 7. RAW_MATERIAL__C
# ============================================================
doc.add_heading('Raw_Material__c — Raw Material', 1)
doc.add_paragraph('Ingredient catalogue used in cosmetic recipes.')
add_table(doc,
    ['API Name', 'Label', 'Type', 'Description'],
    [
        ['Name',                 'Raw Material Name', 'Text',    'Common name of the ingredient'],
        ['INCI__c',              'INCI',              'Text',    'Official name per INCI nomenclature (International Nomenclature of Cosmetic Ingredients)'],
        ['Material_Category__c', 'Material Category', 'Picklist','Tensionactivo · Principio activo · Aceite esencial/Aroma · Liquido · Conservante · Emulsionante · Aceites'],
        ['Quantity__c',          'Quantity',          'Number',  'Stock in grams'],
        ['Price__c',             'Price',             'Currency','Price in euros'],
    ],
    col_widths=[4.5, 4, 3.5, 6.5]
)
doc.add_paragraph()

# --- Footer note ---
doc.add_paragraph()
p = doc.add_paragraph('Generated from Salesforce metadata · Alma Natural Cosmetics · 2026')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(8)
p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

out = os.environ.get('DATA_DICT_OUTPUT', os.path.join(os.path.dirname(__file__), 'Alma_Data_Dictionary.docx'))
doc.save(out)
print(f'Saved: {out}')
